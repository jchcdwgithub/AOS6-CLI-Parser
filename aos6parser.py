from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.worksheet.table import Table
from openpyxl.styles import PatternFill, Font
from openpyxl.utils import get_column_letter
import re
import math

cli_to_api_dict = {'ip access-list session' : 'Session ACL',
                   'user-role' : 'Role',
                   'aaa authentication-server radius' : { 'aaa authentication-server radius' :'Radius Server Name',
                                                          'acctport' : 'Rad Server Acctport',
                                                          'authport' : 'Rad Server Authport',
                                                          'use-md5' : 'Rad Server MD5',
                                                          'enable-radsec' : 'Radsec',
                                                          'radsec-client-cert-name' : 'Radsec Client Cert Name',
                                                          'radsec-port' : 'Radsec Port',
                                                          'radsec-trusted-cacert-name' : 'Radsec Trusted CA Cert Name',
                                                          'radsec-trusted-servercert-name' : 'Radsec Trusted Server Cert Name',
                                                          'retransmit' : 'Rad Server Retransmit',
                                                          'service-type-framed-user' : 'Service Type Framed User',
                                                          'timeout' : 'Rad Server Timeout',
                                                          'use-ip-for-calling-station' : 'Use IP For Calling Station',
                                                          'enable-ipv6' : 'Enable IPv6',
                                                          'host' : 'Rad Server Host',
                                                          'called-station-id type' : 'Called-station-ID Type',
                                                          'called-station-id delimiter' : 'Called-station-ID Delimiter',
                                                          'called-station-id include-ssid' : 'Called-station-ID Include SSID',
                                                          'cppm username' : 'CPPM Username',
                                                          'cppm password' : 'CPPM PW',
                                                          'source-interface vlan' : 'Source Interface VLAN',
                                                          'source-interface ip6addr' : 'Source Interface IPv6 Addr',
                                                          'key' : 'Rad Server Key',
                                                          'nas-identifier' : 'Rad Server NAS-ID',
                                                        },
                    'aaa server-group' : { 'aaa server-group' : 'Server Group',
                                           'auth-server' : 'SG Server Name',},
                    'aaa profile' : { 'aaa profile': 'AAA Profile',
                                      'rfc-3576-server' : 'AAA RFC3576 IP',
                                      'dot1x-default-role' : 'AAA 1X Default Role',
                                      'dot1x-server-group': 'AAA 1X Server Group',
                                      'authentication-dot1x' : 'AAA 1X Profile',
                                      'authentication-mac' : 'AAA MAC Profile',
                                      'initial-role' : 'AAA Initial Role',
                                      'mac-default-role' : 'AAA MAC Default Role',
                                      'mac-server-group' : 'AAA MAC Server Group',
                                      'download-role': 'AAA DL Role',
                                      'l2-auth-fail-through': 'AAA L2 Auth Failthrough',},
                    'ap system-profile' : { 'ap system-profile' : 'AP System Profile',
                                            'lms-ip' : 'AP Sys LMS IP',
                                            'bkup-lms-ip' : 'AP Sys Bkup LMS IP',
                                            'shell-passwd' : 'AP Sys AP Console PW'},
                    'ap regulatory-domain-profile' : { 'ap regulatory-domain-profile' : 'Reg Domain Profile',
                                                       'valid-11a-channel' : '5 GHz Channels',
                                                       'valid-11g-channel' : '2.4 GHz Channels',
                                                       'valid-11a-40mhz-channel-pair' : '5 GHz 40 MHz Channels',
                                                       'valid-11g-40mhz-channel-pair' : '2.4 GHz 40 MHz Channels',
                                                       'valid-11a-80mhz-channel-group': '5 GHz 80 MHz Channels',
                                                       'country-code' : 'Reg Domain Country Code',
                                                       },
                    'rf arm-profile' : { 'rf arm-profile' : 'RF ARM Profile',
                                         '40MHz-allowed-bands' : '40 MHz Allowed Bands',
                                         '80MHz support' : '80 MHz Support',
                                         'acceptable-coverage-index' : 'Acceptable Coverage Index',
                                         'active-scan' : 'Active Scan',
                                         'aggressive-scan' : 'Aggressive Scan',
                                         'assignment' : 'Assignment',
                                         'backoff-time' : 'ARM Backoff Time',
                                         'cellular-handoff-assist' : 'Cellular Handoff Assist',
                                         'channel-quality-aware-arm' : 'Channel Quality Aware ARM',
                                         'channel-quality-aware-threshold' : 'Channel Quality Aware Threshold',
                                         'channel-quality-wait-time' : 'Channel Quality Wait Time',
                                         'client-aware' : 'ARM Client Aware',
                                         'client-match' : 'ARM Client Match',
                                         'cm-band-a-min-signal' : 'CM Band A Min Sig',
                                         'cm-band-g-max-signal' : 'CM Band G Max Sig',
                                         'cm-lb-client-thresh' : 'CM LB Client Threshold',
                                         'cm-lb-signal-delta' : 'CM LB Sig Delta',
                                         'cm-lb-thresh' : 'CM LB Threshold',
                                         'cm-max-steer-fails' : 'CM Max Steer Fails',
                                         'cm-mu-client-thresh' : 'CM MU Client Threshold',
                                         'cm-mu-snr-thresh' : 'CM MU SNR Threshold',
                                         'cm-report-interval' : 'CM Report Interval',
                                         'cm-stale-age' : 'CM Stale Age',
                                         'cm-steer-timeout' : 'CM Steer Timeout',
                                         'cm-sticky-check_intvl' : 'CM Sticky Check Interval',
                                         'cm-sticky-min-signal' : 'CM Sticky Min Signal',
                                         'cm-sticky-snr' : 'CM Sticky SNR',
                                         'cm-sticky-snr-delta' : 'CM Sticky SNR Delta',
                                         'cm-update-interval' : 'CM Update Interval',
                                         'error-rate-threshold' : 'Error Rate Threshold',
                                         'error-rate-wait-time' : 'Error Rate Wait Time',
                                         'free-channel-index' : 'Free Channel Index',
                                         'ideal-coverage-index' : 'Ideal Coverage Index',
                                         'load-aware-scan-threshold' : 'Load Aware Scan Threshold',
                                         'max-tx-power' : 'Max TX Power',
                                         'min-scan-time' : 'Min Scan Time',
                                         'min-tx-power' : 'Min TX Power',
                                         'mode-aware' : 'Mode Aware',
                                         'multi-band-scan' : 'Multi-Band Scan',
                                         'ota-updates' : 'OTA Updates',
                                         'ps-aware-scan' : 'PS Aware Scan',
                                         'rogue-ap-aware' : 'Rogue AP Aware',
                                         'scan mode' : 'ARM Scan Mode',
                                         'scan-interval' : 'ARM Scan Interval',
                                         'scanning' : 'ARM Scanning',
                                         'video-aware-scan' : 'ARM Video Aware Scan',
                                         'voip-aware-scan' : 'ARM VOIP Aware Scan',

                    },
                    'rf optimization-profile' : { 'rf optimization-profile' : 'RF Opt Profile',
                                                  'handoff-assist' : 'Handoff Assist',
                                                  'low-rssi-threshold' : 'Low RSSI Threshold',
                                                  'rssi-check-frequency' : 'RSSI Check Freq',
                                                  'rssi-falloff-wait-time' : 'RSSI Falloff Wait Time',

                    },
                    'rf event-thresholds-profile' : { 'rf event-thresholds-profile' : 'RF Event Thresholds Profile',
                                                      'bwr-high-wm' : 'BWR High WM',
                                                      'bwr-low-wm' : 'BWR Low WM',
                                                      'detect-frame-rate-anomalies' : 'Detect Frame Rate Anomalies',
                                                      'fer-high-wm' : 'FER High WM',
                                                      'fer-low-wm' : 'FER Low WM',
                                                      'frr-high-wm' : 'FRR High WM',
                                                      'frr-low-wm' : 'FRR Low WM',
                                                      'flsr-high-wm' : 'FLSR High WM',
                                                      'flsr-low-wm' : 'FLSR Low WM',
                                                      'fnur-high-wm' : 'FNUR High WM',
                                                      'fnur-low-wm' : 'FNUR Low WM',
                                                      'frer-high-wm' : 'FRER High WM',
                                                      'frer-low-wm' : 'FRER Low WM',
                    },
                    'rf am-scan-profile' : { 'rf am-scan-profile' : 'RF AM Scan Prof',
                                             'dwell-time-active-channel' : 'Dwell Time Active Chan',
                                             'dwell-time-other-reg-domain-channel' : 'Dwell Time Other Reg Dom Chan',
                                             'dwell-time-reg-domain-channel' : 'Dwell Time Reg Dom Chan',
                                             'scan-mode' : 'AM Scan Prof Scan Mode',

                    },
                    'rf dot11a-radio-profile' : {'rf dot11a-radio-profile' : '5 GHz Radio Prof',
                                                 'am-scan-profile' : '5 GHz AM Scan Prof',
                                                 'arm-profile' : '5 GHz ARM Prof',
                                                 'beacon-period' : '5 GHz Beacon Period',
                                                 'cap-reg-eirp' : '5 GHz Cap Reg EIRP',
                                                 'cell-size-reduction' : 'Cell Size Reduction',
                                                 'channel' : '5 GHz Prof Channel',
                                                 'channel-reuse' : '5 GHz Prof Chan Reuse',
                                                 'csa-count' : 'CSA Count',
                                                 'ht-radio-profile' : 'HT Radio Profile',
                                                 'maximum-distance' : 'Max Distance',
                                                 'mgmt-frame-throttle-interval' : 'MGMT Frame Throttle Interval',
                                                 'mgmt-frame-throttle-limit' : 'MGMT Frame Throttle Limit',
                                                 'mode' : '5 GHz Radio Mode',
                                                 'slb-mode' : 'SLB Mode',
                                                 'slb-update-interval' : 'SLB Update Interval',
                                                 'tpc-power' : 'TPC Power',
                                                 'tx-power' : '5 GHz Radio TX Power',
                    },
                    'rf dot11g-radio-profile' : {'rf dot11g-radio-profile' : '11g Radio Prof',
                                                 'am-scan-profile' : '11g AM Scan Prof',
                                                 'arm-profile' : '11g ARM Prof',
                                                 'beacon-period' : '11g Beacon Period',
                                                 'cap-reg-eirp' : '11g Cap Reg EIRP',
                                                 'cell-size-reduction' : 'Cell Size Reduction',
                                                 'channel' : '11g Prof Channel',
                                                 'channel-reuse' : '11g Prof Chan Reuse',
                                                 'csa-count' : 'CSA Count',
                                                 'ht-radio-profile' : 'HT Radio Profile',
                                                 'maximum-distance' : 'Max Distance',
                                                 'mgmt-frame-throttle-interval' : 'MGMT Frame Throttle Interval',
                                                 'mgmt-frame-throttle-limit' : 'MGMT Frame Throttle Limit',
                                                 'mode' : '11g Radio Mode',
                                                 'slb-mode' : 'SLB Mode',
                                                 'slb-update-interval' : 'SLB Update Interval',
                                                 'tpc-power' : 'TPC Power',
                                                 'tx-power' : '11g Radio TX Power',
                                                 },
                    'wlan ht-ssid-profile' : {'wlan ht-ssid-profile' : 'HT SSID Profile',
                                              '40MHz-enable' : '40 MHz Enable',
                                              '80MHz-enable' : '80 MHz Enable',
                                              'ba-amsdu-enable' : 'BA AMSDU Enable',
                                              'high-throughput-enable' : 'HT enable',
                                              'ldpc' : 'LDPC',
                                              'legacy-stations' : 'Legacy Stations',
                                              'max-rx-ampdu-size' : 'Max RX AMPDU Size',
                                              'max-tx-ampdu-size' : 'Max TX AMPDU Size',
                                              'max-tx-ampdu-count-be' : 'Max TX AMPDU Count BE',
                                              'max-tx-ampdu-count-bk' : 'Max TX AMPDU Count BK',
                                              'max-tx-ampdu-count-vi' : 'Max TX AMPDU Count VI',
                                              'max-tx-ampdu-count-vo' : 'Max TX AMPDU Count VO',
                                              'max-vht-mpdu-size' : 'Max VHT MPDU Size',
                                              'min-mpdu-start-spacing' : 'Min MPDU Start Spacing',
                                              'mpdu-agg' : 'MPDU AGG',
                                              'short-guard-intvl-20MHz' : 'Short GI 20 MHz',
                                              'short-guard-intvl-40MHz' : 'Short GI 40 MHz',
                                              'short-guard-intvl-80MHz' : 'Short GI 80 MHz',
                                              'stbc-rx-streams' : 'STBC RX Streams',
                                              'stbc-tx-streams' : 'STBC TX Streams',
                                              'supported-mcs-set' : 'Supported MCS Set',
                                              'temporal-diversity' : 'Temporal Diversity',
                                              'very-high-throughput-enable' : 'VHT Enable',
                                              'vht-mu-txbf-enable' : 'VHT MU TXBF Enable',
                                              'vht-supported-mcs-map' : 'VHT Supported MCS Map',
                                              'vht-txbf-explicit-enable' : 'VHT TXBF Exp Enable',
                                              'vht-txbf-sounding-interval' : 'VHT TXBF Sounding Interval'
                                              },
                    'wlan ssid-profile' : {'wlan ssid-profile' : 'SSID Profile',
                                           'a-basic-rates' : 'A Rates Required',
                                           'a-tx-rates' : 'A Rates Allowed',
                                           'ageout' : 'Station Ageout',
                                           'auth-req-thresh' : 'Auth Request Threshold',
                                           'dtim-period' : 'DTIM Period',
                                           'edca-parameters-profile ap' : 'EDCA Profile AP',
                                           'edca-parameters-profile station' : 'EDCA Profile Client STA',
                                           'essid' : 'ESSID',
                                           'g-basic-rates' : 'G Rates Required',
                                           'g-tx-rates' : 'G Rates Allowed',
                                           'ht-ssid-profile' : 'HT SSID Prof',
                                           'max-clients' : 'SSID Max Clients',
                                           'max-retries' : 'SSID Max Retries',
                                           'max-tx-fails' : 'SSID Max TX Fails',
                                           'opmode' : 'WLAN OPMODE',
                                           'rts-threshold' : 'RTS Threshold',
                                           'wepkey1' : 'SSID WEPKEY1',
                                           'wepkey2' : 'SSID WEPKEY2',
                                           'wepkey3' : 'SSID WEPKEY3',
                                           'wepkey4' : 'SSID WEPKEY4',
                                           'weptxkey' : 'SSID WEP TX KEY',
                                           'wmm-be-dscp' : 'WMM BE DSCP',
                                           'wmm-bk-dscp' : 'WMM BK DSCP',
                                           'wmm-ts-min-inact-int' : 'WMM TS MIN Inactivty Interval',
                                           'wmm-vi-dscp' : 'WMM VI DSCP',
                                           'wmm-vo-dscp' : 'WMM VO DSCP',
                                           'wpa-hexkey' : 'WPA Hexkey',
                                           'wpa-passphrase' : 'WPA PW',
                                           },
                   
                   }
object_values_dict = {}                    
one_line_objects_dict = {}
multi_value_parameter = {}
aos6_cli_rules = {}

def parse():
    try:
        with open('parser_test.txt') as config_file:
            objects = []
            lines = config_file.readlines()
            current_line = 0
            while current_line < len(lines):
                current_object_lines = []
                current_object_lines.append(lines[current_line].strip())
                current_line += 1
                while '!\n' != lines[current_line]: 
                    sub_line = lines[current_line].strip()
                    if sub_line != '':
                        current_object_lines.append(sub_line)
                    current_line += 1
                    if current_line == len(lines):
                        break
                current_line += 1
                objects.append(current_object_lines)
            ir_doc = Document()
            ir_doc.save('acl_parse.docx')
    except FileNotFoundError:
        print('file does not exist.')
        exit()

def gather_objects(lines):
    
    current = 0
    objects = {}
    while current < len(lines):
        current_line = lines[current]
        current_object = [current_line]
        current_object_name = get_object_name_from_line(current_line)
        current += 1
        if current == len(lines):
            break
        while lines[current].strip() != '!':
            current_object.append(lines[current])
            current += 1
            if current == len(lines):
                break
        if current_object_name in objects:
            objects[current_object_name].append(current_object)
        else:
            objects[current_object_name] = [current_object]
        current += 1
    return objects

def process_objects(objects):

    objects_attributes = {} 
    for object_name in objects:
        object_list = objects[object_name]
        object_attributes = []
        for object in object_list:
            object_attributes.append(group_object_attributes(object))
        objects_attributes[object_name] = object_attributes
    return objects_attributes 

def build_attributes_arrays(objects_attributes):

    attributes_arrays = []
    for objects in objects_attributes:
        attributes_object = add_keys_to_attributes_object(objects_attributes[objects]) 
        for object_attributes in objects_attributes[objects]:
            for attribute in object_attributes:
                processed_attribute_list = process_attribute_list(object_attributes[attribute])
                attributes_object[attribute].append(processed_attribute_list)
            for attribute in attributes_object:
                if not attribute in object_attributes:
                    attributes_object[attribute].append('')
        attributes_arrays.append(attributes_object)
    return attributes_arrays
            
def add_keys_to_attributes_object(objects):
    attributes_object = {}
    for object in objects:
        for key in object:
            if not key in attributes_object:
                attributes_object[key] = []
    return attributes_object

def process_attribute_list(list):

    if len(list) == 0:
        return ''
    processed_list = list[0]
    for word in list[1:]:
        processed_list += f' ,{word}'
    return processed_list

def get_object_name_from_line(object):
    words = object.strip().split(' ')
    return join_words(words[:-1], ' ') 

def group_object_attributes(object):
    groups = {}
    name_line = join_words(object[0].split(' ')[:-1], ' ')
    input_name = object[0].strip().split(' ')[-1].replace('"','')
    object_name = name_line
    groups[object_name] = [input_name]

    if object_name in object_values_dict:    
        for line in object[1:]:
            line = line.strip().replace('"', '')
            line_words = line.split(' ')
            is_processed = False
            for unique_object in object_values_dict[object_name]:
                unique_values = unique_object.split(' ')
                match = 0
                current_group_name = ''
                for unique_value,word in zip(unique_values,line_words):
                    if word == unique_value or word in unique_value:
                        match += 1
                        if current_group_name == '':
                            current_group_name += word
                        else:
                            current_group_name += ' ' + word
                        if match == len(unique_values):
                            is_processed = True
                            if current_group_name in groups:
                                groups[current_group_name].append(line[len(current_group_name)+1:])
                            else:
                                groups[current_group_name] = [line[len(current_group_name)+1:]]
                if is_processed:
                    break
    elif object_name in special_objects_dict:
        groups = special_objects_dict[object_name](object[1:], groups)
    return groups

def get_max_columns(objects):
    max_columns = 0
    max_object = None 
    for object in objects:
        current_object_len = len(object)
        if current_object_len > max_columns:
            max_columns = current_object_len
            max_object = object
    return [max_columns,max_object]

def add_acl_info_to_document(doc, acl_objects):
    table = doc.add_table(cols=2,rows=len(acl_objects))
    rows = table.rows
    for row,object in zip(rows,acl_objects):
        row_cells = row.cells
        row_cells[0].text = split_object_identifier(object[0].split(' '), 3)
        row_cells[1].text = object[1]
        for ace in object[2:]:
            row_cells[1].text += f",\n{ace}"

def split_object_identifier(words_list, start):
    object_id = words_list[start:]
    object_id = [word.replace('"','') for word in object_id]
    return join_words(object_id, '_')

def join_words(word_list, seperator):
    if len(word_list) == 0:
        return ''
    joined_words = word_list[0]
    for word in word_list[1:]:
        joined_words += seperator + word
    return joined_words

def process_expression(line, current, origin_id):
    current_id = origin_id
    options = []
    new_options = []
    optional_index = 0
    while current < len(line):
        current_item = line[current]
        if current_item == '|':
            current += 1
            new_options = add_options([current_id],new_options)
            current_id = origin_id
        elif current_item == '{':
            current += 1
            if len(new_options) == 0:
                current, new_options = process_expression(line, current, current_id)
            else:
                current, added_opts = process_expression(line, current, current_id)
                new_options = add_options(new_options, added_opts)
        elif current_item == '}':
            current += 1
            if current == len(line):
                return [current, options]
            deep_current = current
            added_options = []
            if not current_id in new_options:
                new_options.append(current_id)
            options = add_options(options,new_options)
            if len(options) != 0 and line[current] != '|': 
                for option in options:
                    current, extended_options = process_expression(line, deep_current, option)
                    added_options += extended_options
            options = add_options(options,added_options)
            return [current, options]
        elif current_item == '[':
            if not current_id in new_options:
                new_options.append(current_id)
            current += 1
            for index, word in enumerate(line[current:]):
                if word == ']':
                    optional_index = current + index + 1
                    break
            _, extended_options = process_expression(line, optional_index, current_id)
            new_options = add_options(new_options, extended_options)
        elif current_item == ']':
            if not current_id in new_options:
                new_options.append(current_id)
            current += 1
        else:
            current_id += f' {line[current]}'
            current += 1
            if current == len(line):
                options.append(current_id)
    if len(options) == 0:
        options = new_options
    return [current, options]

def add_options(options, added_options):
    if len(options) == 0:
        return added_options
    is_subset = True
    for option in options:
        is_not_in_any = True 
        for add_option in added_options:
            if option in add_option:
                is_not_in_any = False
                break
        if is_not_in_any:
            is_subset = False
            break
    if is_subset:
        options = added_options
    else:
        for added_option in added_options:
            if not added_option in options:
                options.append(added_option)
    return options

def process_line(line):
    current = 1
    words = line.strip().split(' ')
    options = process_expression(words, current, words[0])
    return options[1]

def create_expanded_cli_commands():
    cli_commands_path = 'test_cli_commands.txt'
    expanded_cli_commands = 'more_cli_commands.txt'
    cli_lines = []
    with open(cli_commands_path) as cli_file:
        cli_lines = cli_file.readlines()
    with open(expanded_cli_commands, 'w') as expanded_file:
        indentation = '    '
        for line in cli_lines:
            if line == '!\n' or line == '!':
                expanded_file.write(line)
            else:
                if ' ' != line[0]:
                    indentation = ''
                processed_line = process_line(line)
                for new_line in processed_line:
                    expanded_file.write(f'{indentation}{new_line}\n')

def populate_objects_value_dict():
    cli_commands_path = 'expanded_cli_commands.txt'
    with open(cli_commands_path) as cli_file:
        user_input = re.compile(r'<.+>')
        lines = cli_file.readlines()
        current = 0
        while(current < len(lines)):
            if user_input.match(lines[current].split(' ')[-1]):
                object_name = join_words(lines[current].strip().split(' ')[:-1], ' ')
                object_values_dict[object_name] = []
                current += 1
                while(current < len(lines) and lines[current] != '!\n'):
                    current_line = lines[current].strip()
                    if is_simple_parameter(current_line):
                        keywords = current_line.split(' ')
                        if len(keywords) > 1:
                            keywords = keywords[:-1]
                        command = join_words(keywords, ' ')
                        object_values_dict[object_name].append(command)
                    else:
                        if object_name in multi_value_parameter:
                            multi_value_parameter[object_name].append(current_line.split(' ')[0])
                        else:
                            multi_value_parameter[object_name] = [current_line.split(' ')[0]]
                    current += 1
                current += 1

def is_simple_parameter(line):
    words = line.split(' ')
    is_simple = True
    multiple_user_input = False
    user_input = re.compile(r'<\w+>')
    keyword_pairs = {}
    current_keyword = ''
    for word in words:
        if user_input.match(word):
            if current_keyword in keyword_pairs:
                is_simple = False
                multiple_user_input = True
                break
            else:
                keyword_pairs[current_keyword] = [word]
        else:
            if current_keyword in keyword_pairs:
                current_keyword = '' 
            if current_keyword == '':
                if len(keyword_pairs.keys()) == 1:
                    is_simple = False
                    break
                else:
                    current_keyword = word
            else:
                current_keyword += f' {word}'
    if len(keyword_pairs.keys()) > 1 or multiple_user_input:
        is_simple = False
    return is_simple

def process_bandwidth_contract(bwc_line):
    bwc_words = bwc_line.strip().split(' ')
    bwc_dict = {}
    base_name = bwc_words[0]
    current = 0
    current_word = bwc_words[current]
    if current_word == 'app' or current_word == 'appcategory':
        current_name = base_name + ' ' + 'application'
        bwc_dict[current_name] = current_word
        current_name = base_name + ' ' + current_word
        current += 1
        bwc_dict[current_name] = bwc_words[current]
        current += 1
        bwc_dict[base_name] = bwc_words[current]
        current += 1
        current_name = current_name + ' ' + bwc_words[current]
        bwc_dict[current_name] = bwc_words[current]
    else:
        bwc_dict[base_name] = bwc_words[current]
        current += 1
        current_name = base_name + ' direction'
        bwc_dict[current_name] = bwc_words[current]
    return bwc_dict

def process_int_access_group(ag_line):
    ag_words = ag_line.strip().split(' ')
    ag_dict = {}
    base_name = ag_words[0] + ' ' + ag_words[1]
    current = 2
    ag_name = ag_words[current]
    current += 1
    current_word = ag_words[current]
    name = base_name + ' ' + current_word
    ag_dict[name] = ag_name
    return ag_dict

def is_in_multi_value_dict(object, line):
    word = line.strip().split(' ')[0]
    return word in multi_value_parameter[object]

def process_acl(aces, group):
    joined_aces = []
    for ace in aces:
        joined_aces.append(ace.strip())
    group['acl session aces'] = joined_aces
    return group

def process_bwc(line):
    bwc_values = []
    line_words = line.split(' ')
    header = f"{line_words[0]} {line_words[1]}"
    bwc_name = ''
    if '"' in line_words[2]:
        bwc_name += line_words[2].replace('"', '')
        current = 3
        while current < len(line_words):
            if '"' in line_words[current]:
                bwc_name += '_' + line_words[current].replace('"','')
                current += 1
                break
            else:
                bwc_name += '_' + line_words[current]
                current += 1
    bwc_name_header = f"{header} {bwc_name}"
    bwc_values.append({bwc_name_header: bwc_name})
    bwc_type = f"{header} {line_words[current]}"
    bwc_values.append({bwc_type: line_words[current+1]})
    return bwc_values

def transform_attribute_array_to_array_tables(object_name, attribute_array):
    table_headers = []
    #object_dict = cli_to_api_dict[object_name]
    arrays = []
    for key in attribute_array:
        table_headers.append(key)
        arrays.append(attribute_array[key])
    number_of_objects = len(arrays[0])
    current = 0
    zipped_arrays = [table_headers]
    while current < number_of_objects:
        current_array = []
        for array in arrays:
            current_array.append(array[current])
        current += 1
        zipped_arrays.append(current_array)
    return zipped_arrays


def build_tables_arrays(attributes_array):
    tables_arrays = []
    for object in attributes_array:
        object_name = list(object.keys())[0]
        tables_arrays.append(transform_attribute_array_to_array_tables(object_name,object))
    return tables_arrays

def write_tables_to_excel_worksheets(tables_arrays,output_file):
    grouped_tables = group_tables(tables_arrays)
    workbook = Workbook()
    first_worksheet = True
    for group in grouped_tables:
        if first_worksheet:
            current_worksheet = workbook.active
            current_worksheet.title = group[0]
            first_worksheet = False
        else:
            current_worksheet = workbook.create_sheet(title=group[0])
        current_row = 1
        for table in group[1]:
            table_name = table[0][0].replace(' ','_')
            table = add_node_column_to_table(table)
            current_row = add_table_to_worksheet(table, current_worksheet, table_name=table_name, start=current_row) 
        workbook.save(output_file)

def write_show_tables_to_excel_worksheets(tables_arrays,output_file='show_output.xlsx',workbook=''):
    """ Write the tables in the tables_arrays to the workbook, if given. """

    if workbook == '':
        workbook = Workbook()
    first_worksheet = True
    for show_command in tables_arrays:
        if first_worksheet:
            current_worksheet = workbook.active
            current_worksheet.title = show_command
            first_worksheet = False
        else:
            if len(show_command) > 30:
                title = show_command[:30]
            else:
                title = show_command
            current_worksheet = workbook.create_sheet(title=title)
        table_name = show_command
        add_table_to_worksheet(tables_arrays[show_command], current_worksheet, table_name=table_name, start=1)
    workbook.save(output_file)

def add_node_column_to_table(table):
    ''' Adds a node column to the table. '''

    table[0].insert(0,'Node')
    for row in table[1:]:
        row.insert(0,'')
    
    return table

def group_tables(tables_arrays):
    grouped_tables = {}
    for table in tables_arrays:
        group_membership = classify_table(table[0][0])
        if group_membership in grouped_tables:
            grouped_tables[group_membership].append(table)
        else:
            grouped_tables[group_membership] = [table]
    grouped_tables_pairs = []
    for membership in grouped_tables:
        grouped_tables_pairs.append((membership,grouped_tables[membership]))
    return grouped_tables_pairs

def classify_table(table_header):
    if is_in_ap_ws(table_header):
        return 'ap'
    elif is_in_roles_ws(table_header):
        return 'roles_policies'
    elif is_in_wlan_ws(table_header):
        return 'wlan'
    elif is_in_security_ws(table_header):
        return 'security'
    elif is_in_net_services_ws(table_header):
        return 'net_services'
    elif is_in_interfaces_ws(table_header):
        return 'interface'
    elif is_in_radio_ws(table_header):
        return 'radio'
    elif is_in_crypto_ws(table_header):
        return 'crypto'
    elif is_in_controller_ws(table_header):
        return 'controller'
    elif is_in_ids_ws(table_header):
        return 'ids_wms'
    elif is_in_services_ws(table_header):
        return 'services'
    elif is_in_l2_l3_ws(table_header):
        return 'l2_l3'
    else:
        return 'misc'

def is_in_ap_ws(table_header):
    starts_with_ap = re.compile(r'^ap[ -]')
    starts_with_iap = re.compile(r'^iap')
    starts_with_provision_ap = re.compile(r'^provision-ap')
    return is_in_worksheet([starts_with_ap, starts_with_iap, starts_with_provision_ap], table_header)

def is_in_l2_l3_ws(table_header):
    starts_with_ip = re.compile(r'^ip')
    starts_with_spt = re.compile(r'^spanning-tree')
    starts_with_router = re.compile(r'^router')
    starts_with_vrrp = re.compile(r'^vrrp')
    starts_with_vlan = re.compile(r'^vlan')
    return is_in_worksheet([starts_with_ip, starts_with_spt, starts_with_router, starts_with_vrrp, starts_with_vlan], table_header)

def is_in_wlan_ws(table_header):
    starts_with_wlan = re.compile(r'^wlan')
    return is_in_worksheet([starts_with_wlan], table_header)

def is_in_security_ws(table_header):
    starts_with_aaa = re.compile(r'^aaa')
    starts_with_local = re.compile(r'^local')
    return is_in_worksheet([starts_with_aaa, starts_with_local], table_header)

def is_in_net_services_ws(table_header):
    starts_with_ntp = re.compile(r'^ntp')
    starts_with_clock = re.compile(r'^clock')
    starts_with_ip_name = re.compile(r'^ip name')
    starts_with_ip_dhcp = re.compile(r'^ip dhcp')
    starts_with_snmp = re.compile(r'^snmp')
    return is_in_worksheet([starts_with_ntp, starts_with_clock, starts_with_ip_dhcp, starts_with_ip_name, starts_with_snmp], table_header)

def is_in_interfaces_ws(table_header):
    starts_with_interface = re.compile(r'^interface')
    starts_with_lacp = re.compile(r'^lacp')
    return is_in_worksheet([starts_with_interface, starts_with_lacp], table_header)

def is_in_radio_ws(table_header):
    starts_with_rf = re.compile(r'^rf ')
    return is_in_worksheet([starts_with_rf], table_header)

def is_in_crypto_ws(table_header):
    starts_with_crypto = re.compile(r'^crypto')
    starts_with_tunnel = re.compile(r'tunnel')
    starts_with_vpdn = re.compile(r'^vpdn')
    return is_in_worksheet([starts_with_crypto, starts_with_tunnel, starts_with_vpdn], table_header)

def is_in_controller_ws(table_header):
    starts_with_controller = re.compile(r'^controller')
    starts_with_cluster = re.compile(r'^cluster')
    starts_with_master = re.compile(r'^master')
    starts_with_whitelistdb = re.compile(r'^whitelist-db')
    starts_with_upgrade = re.compile(r'^upgrade')
    starts_with_license = re.compile(r'^license')
    return is_in_worksheet([starts_with_cluster, starts_with_controller, starts_with_master, starts_with_whitelistdb, starts_with_upgrade, starts_with_license], table_header)

def is_in_services_ws(table_header):
    starts_with_airgroup = re.compile(r'^airgroup')
    starts_with_app = re.compile(r'^app')
    starts_with_esi = re.compile(r'^esi')
    starts_with_voice = re.compile(r'^voice')
    starts_with_pan = re.compile(r'^pan')
    starts_with_web_cc = re.compile(r'^web-cc')
    return is_in_worksheet([starts_with_airgroup, starts_with_app, starts_with_esi, starts_with_voice, starts_with_pan, starts_with_web_cc], table_header)

def is_in_roles_ws(table_header):
    starts_with_user_role = re.compile(r'^user-role')
    starts_with_ip_access = re.compile(r'^ip access-list')
    starts_with_netdest = re.compile(r'^netdest')
    return is_in_worksheet([starts_with_user_role, starts_with_ip_access, starts_with_netdest], table_header)

def is_in_ids_ws(table_header):
    starts_with_ids = re.compile(r'^ids')
    starts_with_wms = re.compile(r'^wms')
    return is_in_worksheet([starts_with_ids, starts_with_wms], table_header)

def is_in_worksheet(list_of_regexes, table_header):
    is_member = False
    for regex in list_of_regexes:
        if regex.match(table_header):
            is_member = True
            break
    return is_member
    
def write_to_excel(title, data, workbook, start):
    
    current_ws = workbook.create_sheet(title=title)
    title_row_color = 'C0504D'

    for row in data:
        current_ws.append(row)

    colors = ['E6B8B7', 'F2DCDB']
    color_index = 0
    rows = current_ws.rows
    for row in rows:
        if color_index == 0:
            for cell in row:
                cell.fill = PatternFill('solid', fgColor = title_row_color)
                cell.font = Font(color='FFFFFF')
        else:
            for cell in row:
                cell.fill = PatternFill('solid', fgColor=colors[color_index%2])
                cell.font = Font(color='000000')
        color_index += 1
    
    column_widths = []
    for row in data:
        for i, cell in enumerate(row):
            if len(column_widths) > i:
                if len(cell) > column_widths[i]:
                    column_widths[i] = len(cell)
            else:
                column_widths += [len(cell)]
    
    for i, column_width in enumerate(column_widths,1):  # ,1 to start at 1
        current_ws.column_dimensions[get_column_letter(i)].width = column_width + 5

    table = Table(displayName='SSID',ref=f'A{start}:F{start + len(data)}')
    current_ws.add_table(table)

def add_table_to_worksheet(data,worksheet,table_name='',start=1):
    
    end_cell_number = len(data)
    end_cell = (end_cell_number, start + len(data[0]))
    start_cell = (1,start)
    num_columns = len(data[0])

    current = 1 
    data_index = 0
    while current <= end_cell_number:
        current_column = start
        current_row_cells = []
        while current_column < start + num_columns:
            current_row_cells.append(worksheet.cell(row=current, column=current_column))
            current_column += 1
        current_data_row = data[data_index]
        for cell,data_value in zip(current_row_cells,current_data_row):
            cell.value = data_value
        current += 1
        data_index += 1 
    add_color_scheme(worksheet,start_cell,end_cell)
    column_widths = get_widest_column_widths(start-1,data)
    adjust_column_widths(worksheet,column_widths)

    return start + len(data[0]) + 3

def adjust_column_widths(worksheet,column_widths):
    ''' Adjust the column widths of the table in the worksheet based on the dictionary passed. '''

    for column_letter in column_widths:
        worksheet.column_dimensions[column_letter].width = column_widths[column_letter] + 5

def add_color_scheme(worksheet,start_cell,end_cell):
    ''' Add a red table color scheme to the cells in the worksheet. '''
    title_row_color = 'C0504D'
    colors = ['E6B8B7', 'F2DCDB']
    color_index = 0
    rows = []
    start_row,start_column = start_cell
    end_row,end_column = end_cell
    current_row = 1
    while current_row <= end_row:
        current_column = start_column 
        current_row_list = []
        while current_column < end_column:
            current_row_list.append(worksheet.cell(row=current_row,column=current_column))
            current_column += 1
        rows.append(current_row_list)
        current_row += 1
    for row in rows:
       if color_index == 0:
          for cell in row:
               cell.fill = PatternFill('solid', fgColor = title_row_color)
               cell.font = Font(color='FFFFFF')
       else:
            for cell in row:
               cell.fill = PatternFill('solid', fgColor=colors[color_index%2])
               cell.font = Font(color='000000')
       color_index += 1

def get_widest_column_widths(start,data):
    ''' Returns a dictionary of column letter to widest length of data found in each column. '''
    column_widths = {}
    columns_table = swap_rows_and_columns(data)
    column_index = start
    for column in columns_table:
        column_letter = calculate_column_letters(column_index)
        current_longest_len = 0
        for cell in column:
            if len(cell) > current_longest_len:
                column_widths[column_letter] = len(cell)
                current_longest_len = len(cell)
        column_index += 1
    return column_widths

def swap_rows_and_columns(data):
    ''' Given an array of arrays, swap the rows and columns data. '''

    num_rows = len(data)
    swapped_matrix = []
    num_cols = len(data[0])
    current_col = 0
    while current_col < num_cols:
        current_row = 0
        current_col_array = []
        while current_row < num_rows:
            current_col_array.append(data[current_row][current_col])
            current_row += 1
        current_col += 1
        swapped_matrix.append(current_col_array)

    return swapped_matrix

def calculate_column_letters(column_index):
    ''' Given a column index, return the corresponding excel column letter. '''

    column_letters = [letter.upper() for letter in 'abcdefghijklmnopqrstuvwxyz']

    if column_index < 26:
        return column_letters[column_index]
    else:
        first_letter_index = math.floor(column_index/26) - 1
        last_letter_index = column_index%26 - 1

        first_letter = column_letters[first_letter_index]
        last_letter = column_letters[last_letter_index]

        column = first_letter + last_letter
        return column
            
def extract_information_from_rule(rule, cli_line):
    ''' Given a rule, return a list of tuples with the rule_header and associated data. '''
    
    header = ''
    data = ''
    pair_list = []
    cli_words = cli_line.split(' ')
    cli_words = remove_empty_values(cli_words)
    rule_words = rule.split(' ')
    first_word = rule_words[0]
    header += first_word
    rule_index = 1
    while rule_index < len(cli_words):
        if not '<' in rule_words[rule_index]:
            header += f" {rule_words[rule_index]}"
            rule_index += 1
        else:
            if header == first_word and len(pair_list) != 0:
                extracted_input_name = rule_words[rule_index].replace('<','').replace('>','')
                header += f" {extracted_input_name}"
            data = cli_words[rule_index]
            if rule_index + 1 < len(cli_words):
                rule_index += 1
                while rule_index < len(cli_words):
                    if '<' in cli_words[rule_index]:
                        data += f" {cli_words[rule_index]}"
                        rule_index += 1
                    else:
                        pair_list.append((header, data))
                        break
            else:
                pair_list.append((header,data))
                rule_index += 1
            header = first_word
            data = ''
    if header.split(' ')[-1] == rule_words[-1] and len(pair_list) > 0:
        current_data = pair_list[-1][1]
        current_header = pair_list[-1][0]
        current_data += f" {rule_words[-1]}"
        pair_list[-1] = (current_header, current_data)
    return pair_list

def remove_empty_values(word_list):
    ''' Removes any empty values form the word list. They represent extra spaces that might have been added to the input. '''

    new_list = []
    for word in word_list:
        if word != '':
            new_list.append(word)
    return new_list

def match_cli_output_to_rule(object_name, cli_output):
    "Given a cli line, match it with the set of rules associated with the object. "
    
    if object_name in aos6_cli_rules:
        cli_rules = aos6_cli_rules[object_name]
        current_index = 0
        correct_rule = ''
        rule_index = 0
        cli_words = cli_output.split(' ')
        cli_words = remove_empty_values(cli_words)
        for rule in cli_rules:
            rule_words = rule.split(' ')
            rule_length = len(rule_words)
            while current_index < len(cli_words) and rule_index < rule_length:
                if '<' in rule_words[rule_index] or cli_words[current_index] == rule_words[rule_index]:
                    current_index += 1
                    rule_index += 1
                else:
                    break
            if current_index == len(cli_words):
                correct_rule = rule
                break
        return correct_rule

def populate_cli_rules(template_file):
    ''' Read from the expanded cli commands and load the rules into the aos6_cli_rules dictionary. '''

    with open(template_file) as cli_txt_rules:
        rules_lines = cli_txt_rules.readlines()
        current_index = 0
        while current_index < len(rules_lines):
            header_name = get_header(rules_lines[current_index])
            aos6_cli_rules[header_name] = []
            while current_index < len(rules_lines):
                current_line = rules_lines[current_index].strip()
                current_index += 1
                if current_line != '!':
                    aos6_cli_rules[header_name].append(current_line)
                else:
                    break

def get_header(rule):
    ''' Returns the name of the object from the first line of the rules section. '''
    
    rule_words = rule.split(' ')
    header = rule_words[0]
    index = 1
    while index < len(rule_words):
        if '<' in rule_words[index]:
            index +=1 
            break
        else:
            header += f" {rule_words[index]}"
            index += 1
    # if more then complicated object name ...
    return header

def group_cli_lines(cli_lines):
    ''' Given a list of cli lines, whitespace and all, group them into the relevant object lines. '''
    
    cli_groups = []
    current = 0
    while current < len(cli_lines):
        current_group = []
        while current < len(cli_lines) and cli_lines[current].strip() != '!':
            current_line = check_special_line(format_names(cli_lines[current].strip()))
            if current_line != '':
                current_group.append(current_line)
            current += 1
        cli_groups.append(current_group)
        current += 1
    return cli_groups

def check_special_line(cli_line):
    ''' Certain cli lines need to be specially pre-processed before returning. '''
    
    special_lines = {'a-basic-rates':group_rates,'g-basic-rates':group_rates, 'a-tx-rates':group_rates, 'g-tx-rates':group_rates}
    if cli_line == '':
        return ''
    else:
        cli_words = cli_line.split(' ')
        cli_title = cli_words[0] 
        if cli_title in special_lines:
            return special_lines[cli_title](cli_line)
        else:
            return cli_line

def group_rates(special_line):
    ''' Rates are comma separated numeric values on the CLI. Group them into a comma separated list and return the new line. '''

    cli_words = special_line.split(' ')
    cli_title = cli_words[0]
    if len(cli_words) > 0:
        current_rates = cli_words[1]
        for rate in cli_words[2:]:
            current_rates += f",{rate}"
        return join_words([cli_title,current_rates],' ')
    else:
        return ''

def make_object_from_cli_group(cli_group):
    ''' Gather data from the group and return an object containing parameter: [data] entries. '''

    formatted_cli_group = format_names(cli_group[0])
    header = formatted_cli_group.split(' ')[:-1]
    header = join_words(header, ' ')
    current_param = []
    current_data = []
    for cli_line in cli_group:
        sanitized_line = format_names(cli_line)
        rule = match_cli_output_to_rule(header, sanitized_line)
        if rule != '' and rule is not None:
            extracted = extract_information_from_rule(rule, sanitized_line)
            for pair in extracted:
                if pair[0] in current_param:
                    pair_index = current_param.index(pair[0])
                    current_data[pair_index].append(pair[1])
                else:
                    current_param.append(pair[0])
                    current_data.append([pair[1]])
    current_object = {}
    for param, data in zip(current_param,current_data):
        if param in current_object:
            current_object[param] += data
        else:
            current_object[param] = data
    return [header, current_object]

def format_names(cli_line):
    ''' Remove double quotes and join names found in double quotes with underscores replacing spaces. '''

    cli_words = cli_line.split(' ')
    replaced_name = ''
    replaced_cli_line = cli_words[0]
    index = 1
    while index < len(cli_words):
        if '"' in cli_words[index]:
            replaced_name += cli_words[index].replace('"','')
            index += 1
            if index == len(cli_words):
                replaced_cli_line += f" {replaced_name}"
            while index < len(cli_words):
                if '"' in cli_words[index]:
                    removed_quotes = cli_words[index].replace('"','')
                    replaced_name += f"_{removed_quotes}"
                    index += 1
                    replaced_cli_line += f" {replaced_name}"
                    replaced_name = ''
                    break
                else:
                    replaced_name += f"_{cli_words[index]}"
                    index += 1
        else:
            replaced_cli_line += f" {cli_words[index]}"
            index += 1
    return replaced_cli_line

def make_cli_objects(cli_file):
    ''' Parses the cli_file and returns a dictionary of key to value entries based on the cli outputs. '''

    with open(cli_file) as f:
        lines = f.readlines()
        cli_groups = group_cli_lines(lines)
        cli_objects = {}
        for cli_group in cli_groups:
            if len(cli_group) > 0:
                header, current_object = make_object_from_cli_group(cli_group)
                if header in cli_objects:
                    cli_objects[header].append(current_object)
                else:
                    if current_object != {}:
                        cli_objects[header] = [current_object]
        return cli_objects

def get_table_range(worksheet, start_column_index):
    ''' Return a range of cells representing a table. Returns the table and the next cell to start the search. '''

    start_cell = (1,start_column_index)
    rows = find_table_end(worksheet,column_index=start_column_index,find_rows=True)
    columns = find_table_end(worksheet,column_index=start_column_index)
    end_row = check_for_added_rows(worksheet, start_column_index, rows, columns)
    end_col = check_for_added_columns(worksheet, columns)
    end_cell = (end_row, end_col)
    return [start_cell,end_cell]

def check_for_added_rows(worksheet, start_column, end_row, end_column):
    ''' Checks whether there have been other data added to the current table and if so will return the updated row count. '''

    rows_added = True
    current_row = end_row+1
    while rows_added:
        columns = [col for col in range(start_column, end_column+1)]
        current_cells = [worksheet.cell(row=current_row, column=col).value for col in columns]
        is_empty = True
        for cell_value in current_cells:
            if cell_value is not None:
                is_empty = False
        if not is_empty:
            current_row += 1
        else:
            rows_added = False
    return current_row - 1

def check_for_added_columns(worksheet, end_column):
    ''' Checks for whether columns have been added to the table and if so will return the updated column count. '''

    columns_added = True
    current_col = end_column+1
    while columns_added:
        potential_header = worksheet.cell(row=1,column=current_col)
        if potential_header.value is not None:
            current_col += 1
        else:
            columns_added = False
    return current_col - 1 

def find_table_end(worksheet,column_index=1,find_rows=False):
    ''' Finds the end of either the column or the row given a worksheet and the column index. The row index always starts at 1.
        Returns either the row index or column index where the table ends. '''
    
    incrementor = increment_rows if find_rows else increment_cols
    current_col = column_index
    current_row = 1
    end_table = False
    while not end_table:
        current_cell = worksheet.cell(row=current_row,column=current_col)
        if current_cell.fill.fgColor.value != '00000000':
            current_row,current_col = incrementor(current_row,current_col)
        else: 
            end_table = True
    return current_row - 1 if find_rows else current_col - 1

def increment_rows(row, col):
    ''' increments the row by one and returns an array of [row, col] '''

    return [row+1,col]

def increment_cols(row, col):
    ''' increments the col by one and returns an array of [row, col] '''

    return [row, col+1]

def find_beginning_of_next_table(worksheet, next_column):
    ''' Returns the index of the next column or -1 if the next table cannot be found at most 3 cells over. '''

    columns_checked = 0
    row = 1
    while columns_checked < 3:
        current_cell = worksheet.cell(row=row,column=next_column)
        current_cell_color = current_cell.fill.fgColor.value
        if current_cell_color != '00000000':
            return next_column
        else:
            next_column += 1
            columns_checked += 1
    return -1

def get_workbook_table_ranges(workbook):
    ''' Returns the table ranges, grouped by worksheets, from the excel workbook. '''

    next_column = 1
    worksheet_dict = {}
    for sheet in workbook:
        worksheet_dict[sheet.title] = []
        while next_column != -1:
            start_cell, end_cell = get_table_range(sheet, next_column)
            worksheet_dict[sheet.title].append([start_cell,end_cell])
            next_column = find_beginning_of_next_table(sheet, end_cell[1]+1)
        next_column = 1 
    return worksheet_dict

def gather_cell_values_into_tables(workbook,tables_dict):
    ''' Gathers all the data in the worksheets into tables, arranged by columns, and returns an array of tables. '''

    tables = []
    for worksheet in tables_dict:
        current_worksheet = workbook[worksheet]
        worksheet_ranges = tables_dict[worksheet]
        for range in worksheet_ranges:
            current_table = []
            start_cell, end_cell = range
            current_column = start_cell[1]
            end_column = end_cell[1]
            end_row = end_cell[0]
            while current_column <= end_column:
                current_row = 1
                column_values = []
                while current_row <= end_row:
                    cell_value = current_worksheet.cell(row=current_row,column=current_column).value
                    if cell_value is None:
                        cell_value = ''
                    column_values.append(cell_value)
                    current_row += 1
                current_table.append(column_values)
                current_column += 1
            tables.append(current_table)
    return tables

def calculate_word_spacing(table_headers):
    ''' Given the table headers, return an array of indices that separates the columns. '''

    dash_index = 0
    table_data_limits = [0]
    while dash_index < len(table_headers):
            #find space
        while dash_index < len(table_headers) and table_headers[dash_index] != ' ':
            dash_index += 1
        while dash_index < len(table_headers) and table_headers[dash_index] == ' ':
            dash_index += 1
        table_data_limits.append(dash_index)
    return table_data_limits

def gather_headers(table_headers,table_headers_underscores):
    """ Given a line of table headers, return a list of headers. """

    header_index = 0
    headers = []
    max_len = len(table_headers) if len(table_headers) < len(table_headers_underscores) else len(table_headers_underscores)
    while header_index < max_len:
        current_header = ''
        while table_headers_underscores[header_index] != ' ':
            current_header += table_headers[header_index]
            header_index += 1
            if header_index == max_len:
                break
        headers.append(current_header)
        while header_index < max_len and table_headers_underscores[header_index] == ' ':
                header_index += 1
    
    return headers

def get_data_lines(log_lines):
    """ Return a list of numbers representing the limits of the different types of show data in the file. """

    current_line = 0
    data_lines = []
    while current_line < len(log_lines):
        if 'show' in log_lines[current_line] and '\n' == log_lines[current_line+1]:
            current_line += 3
            data_lines.append(current_line)
        else:
            current_line += 1
    return data_lines

def group_show_information_into_tables(log_file):
    """ Reads through a file of show commands and returns the information in them in a table. """

    log_lines = log_file.readlines()
    empty_prompt = get_command_line_prompt(log_lines)
    #auto logged files from putty inserts extra \n for whatever reason
    log_lines = clear_excessive_empty_lines(log_lines,empty_prompt)
    data_lines = get_data_lines(log_lines)
    data_lines.append(len(log_lines)-1)
    log_lines,table_names= clear_flag_sections(log_lines,empty_prompt,data_lines)
    tables_groups = group_output_tables(log_lines,table_names)
    tables = {} 
    for table_group in tables_groups:
        extract_information_from_table_group(table_group,tables)
    return tables

def extract_information_from_table_group(table_group,tables):
    """ table_group is a list of tables separated by table titles and table headers. Return a dictionary of the table columns. """

    current_index = 1
    table_title = table_group[0]
    headers = gather_headers(table_group[current_index+1],table_group[current_index+2])
    tables[table_title] = [headers]
    table_positions = index_table_positions(table_group,table_title)
    table_index = 0
    while table_index + 1 < len(table_positions):
        sub_table = table_group[table_positions[table_index]:table_positions[table_index+1]]
        sub_table_rows = gather_table_columns(sub_table)
        table_index += 1
        tables[table_title] += sub_table_rows
    return tables

def gather_table_columns(table):
    """ Given a show table, gather the data from each row into individual column values and return a list of column values. """

    current_index = 0
    table_header_underscores = table[current_index+3]
    table_data_limits = calculate_word_spacing(table_header_underscores)
    current_table = []
    current_index += 4
    while current_index < len(table):
        column_values = []
        data_index = 0
        current_row = table[current_index]
        while data_index+1 < len(table_data_limits):
            start_word = table_data_limits[data_index]
            end_word = table_data_limits[data_index+1]
            current_value = current_row[start_word:end_word]
            column_values.append(current_value)
            data_index += 1
        current_table.append(column_values)
        current_index += 1
    return current_table

def index_table_positions(table_group,table_title):
    """ Return an array of positions for the tables in table_group. """

    table_positions = [] 
    current_index = 0
    while current_index < len(table_group):
        if table_group[current_index] == table_title:
            table_positions.append(current_index)
        current_index += 1
    table_positions.append(len(table_group))
    return table_positions

def clear_excessive_empty_lines(lines,empty_prompt):
    """ Clears any lines containing only newspace. """

    new_lines = []
    current_line = 0
    while current_line < len(lines):
        if lines[current_line] == '\n':
            if empty_prompt in lines[current_line-1]:
                new_lines.append(lines[current_line])
        else:
            new_lines.append(lines[current_line])
        current_line += 1
    return new_lines

def clear_flag_sections(log_lines, empty_prompt, data_lines):
    """ Clears the flag sections between and after tables. """

    data_index = 0
    start_data = data_lines[data_index] - 1
    current_name = log_lines[start_data]
    table_names = [current_name]
    starts_with_flag = re.compile(r'^Flag')
    new_lines = []
    while start_data < len(log_lines):
        if starts_with_flag.match(log_lines[start_data]):
            while log_lines[start_data] != current_name and log_lines[start_data].strip() != empty_prompt and f"{empty_prompt}show" not in log_lines[start_data]:
                start_data += 1 
                if start_data == len(log_lines):
                    break
        elif log_lines[start_data].strip() == empty_prompt:
            data_index += 1
            if data_index < len(data_lines):
                start_data = data_lines[data_index] - 1
                current_name = log_lines[start_data]
                table_names.append(current_name)
            else:
                break
        elif f"{empty_prompt}show" in log_lines[start_data]:
            data_index += 1
            if data_index < len(data_lines):
                start_data = data_lines[data_index] - 1
                current_name = log_lines[start_data]
                table_names.append(current_name)
        else:
            new_lines.append(log_lines[start_data])
            start_data += 1
    return new_lines,table_names

def group_output_tables(log_lines, table_names):
    """ The log lines only contain table information without any of the cli prompts. """

    current_index = 0
    table_index = 1
    if len(table_names) == 1:
        return [log_lines]
    tables = []
    while current_index < len(log_lines):
        current_table = []
        while log_lines[current_index] != table_names[table_index]:
            current_table.append(log_lines[current_index])
            current_index += 1
        tables.append(current_table)
        table_index += 1
        if table_index == len(table_names):
            break
    return tables

def get_command_line_prompt(log_lines):
    """ Returns the command line prompt string. Used to match command line strings later that might appear in a table. """

    for line in log_lines:
        if '#show' in line.replace(' ','') and '(' == line[0]:
            show_index = line.index('#')
            empty_line = line[:show_index+1]
            return empty_line
    return ''

special_objects_dict = {'ip access-list session' : process_acl }